# 文档处理器模块 - 负责加载文档并切分成文本片段
# 【调用者】backend/main.py upload_file 接口
# 【调用外部】LangChain文档加载器、mammoth、textract

# ========== 标准库导入 ==========
import os                                              # 操作系统模块，处理文件路径和环境变量

# ========== LangChain相关导入 ==========
from langchain_text_splitters import RecursiveCharacterTextSplitter  # 递归字符分割器，保持语义完整
from langchain_community.document_loaders import TextLoader, PyPDFLoader, Docx2txtLoader  # 文档加载器
import mammoth                                         # 读取DOCX文件
import textract                                        # 读取旧版DOC文件
from langchain_core.documents import Document          # Document文档对象类型

# ========== 第三方库导入 ==========
from dotenv import load_dotenv                         # 从.env加载环境变量
from typing import List                                # 类型提示模块

# 加载环境变量配置
load_dotenv()


class DocumentProcessor:
    # 【调用者】backend/main.py upload_file 接口调用 process_file()
    # 【字段】chunk_size: 每个片段字符数（默认500，从环境变量读取）
    # 【字段】chunk_overlap: 相邻片段重叠字符数（默认50，保持上下文）
    # 【字段】text_splitter: 文本分割器实例
    
    def __init__(self):
        # 从环境变量读取配置，默认500字符
        self.chunk_size: int = int(os.getenv("CHUNK_SIZE", 500))
        
        # 相邻片段重叠字符数，默认50
        self.chunk_overlap: int = int(os.getenv("CHUNK_OVERLAP", 50))
        
        # 创建递归字符分割器，分隔符优先级：段落>行>空格>字符
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,          # 片段大小
            chunk_overlap=self.chunk_overlap,    # 重叠字符数
            length_function=len,                 # 长度计算方式
            separators=["\n\n", "\n", " ", ""]   # 分隔符优先级
        )
    
    def load_document(self, file_path: str) -> List[Document]:
        # 【调用者】process_file()
        # 【功能】根据文件扩展名选择加载器加载文档
        # 【支持格式】.txt .pdf .docx .doc
        
        ext = os.path.splitext(file_path)[1].lower()  # 获取文件扩展名
        
        try:
            if ext == ".txt":                          # 纯文本文件
                loader = TextLoader(file_path, encoding='utf-8')
            elif ext == ".pdf":                        # PDF文件
                loader = PyPDFLoader(file_path)
            elif ext == ".docx":                       # DOCX文件（可能是ZIP或OLE2格式）
                print(f"【DOCX处理】开始处理文件: {file_path}")
                
                if not os.path.exists(file_path):
                    raise RuntimeError(f"文件不存在: {file_path}")
                
                with open(file_path, 'rb') as f:
                    header = f.read(8)                # 读取文件头判断格式
                
                if header[:2] == b'PK':               # 真正的DOCX（ZIP格式）
                    print(f"【DOCX处理】文件类型: DOCX (ZIP格式)")
                    import mammoth
                    with open(file_path, 'rb') as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        text = result.value
                elif header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':  # DOC格式错误命名为docx
                    print(f"【DOCX处理】文件类型: DOC (OLE2格式)")
                    try:
                        from docx import Document as DocxDocument
                        doc = DocxDocument(file_path)
                        text = '\n'.join([para.text for para in doc.paragraphs])
                    except Exception as e:
                        import olefile
                        try:
                            ole = olefile.OleFileIO(file_path)
                            if ole.exists('WordDocument'):       # 获取Word文档的文本内容
                                stream = ole.openstream('WordDocument')
                                content = stream.read()
                                text = content.decode('utf-16-le', errors='ignore')  # 提取文本
                            else:
                                raise RuntimeError("无法找到WordDocument流")
                        except Exception as ole_error:
                            raise RuntimeError(f"无法处理DOC文件: {str(ole_error)}")
                else:                                           # 默认使用mammoth处理
                    import mammoth
                    with open(file_path, 'rb') as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        text = result.value
                
                documents = [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]
                return documents
            elif ext == ".doc":                                  # 旧版Word文件（OLE2格式）
                import olefile
                try:
                    ole = olefile.OleFileIO(file_path)
                    if ole.exists('WordDocument'):
                        stream = ole.openstream('WordDocument')
                        content = stream.read()
                        text = content.decode('utf-16-le', errors='ignore')
                        text = ''.join([c for c in text if c.isprintable() or c in '\n\r\t'])  # 清理特殊字符
                    else:
                        raise RuntimeError("无法找到WordDocument流")
                except Exception as e:
                    raise RuntimeError(f"无法处理DOC文件: {str(e)}")
                
                documents = [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]
                return documents
            else:
                raise ValueError(f"不支持的文件格式: {ext}")     # 不支持的格式
            
            documents: List[Document] = loader.load()            # 使用加载器读取文档
            return documents
        except ValueError:
            raise                                                # 重新抛出格式错误
        except Exception as e:
            raise RuntimeError(f"加载文档失败: {str(e)}")        # 包装其他异常
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        # 【调用者】process_file()
        # 【功能】将文档切分成小片段，保持语义完整
        split_docs: List[Document] = self.text_splitter.split_documents(documents)
        return split_docs
    
    def process_file(self, file_path: str) -> List[Document]:
        # 【调用者】backend/main.py upload_file()
        # 【功能】处理文件的完整流程：加载→切分→添加元数据
        # 【返回】切分后的Document列表，每个含chunk_id和source元数据
        
        documents = self.load_document(file_path)              # 【调用】load_document()
        split_docs = self.split_documents(documents)          # 【调用】split_documents()
        
        for i, doc in enumerate(split_docs):
            doc.metadata["chunk_id"] = i                      # 添加片段序号
            # source: 原始文件名（不含路径），便于查询时显示来源
            doc.metadata["source"] = os.path.basename(file_path)
        
        return split_docs
    
    # 【调用者】backend/main.py的upload_file()（文件格式校验）
    # 【功能】获取支持的文件格式列表
    # 【返回】[".txt", ".pdf", ".docx", ".doc"]
    def get_supported_extensions(self) -> List[str]:
        return [".txt", ".pdf", ".docx", ".doc"]